from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


doc = Document()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INVENTORY MANAGEMENT SYSTEM (IMS)\n6-MONTH PROGRESS REPORT (ONE-PAGER)")
run.bold = True
run.font.size = Pt(16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Reporting Period: Month 1 to Month 6\nPrepared For: Project Manager").font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "1. Executive Summary", 12)
doc.add_paragraph(
    "Over the last six months, IMS has been developed into a production-ready enterprise platform for core operations. "
    "The implemented scope now covers procurement, purchase orders, delivery receiving, stock management, stock issuance, "
    "multi-level approvals, store keeper verification, and role-based governance."
)

add_heading(doc, "2. Proper Operational Flow (As Implemented)", 12)
add_bullet(doc, "User raises stock issuance request (personal/returnable/individual with item-wise quantities).")
add_bullet(doc, "Wing Supervisor reviews request and performs first-level decision (approve/reject/forward).")
add_bullet(doc, "If physical confirmation is required, item is forwarded to Store Keeper for verification.")
add_bullet(doc, "Store Keeper performs physical count and submits verification: available/partial/unavailable.")
add_bullet(doc, "Approver reviews verification feedback and finalizes approval path.")
add_bullet(doc, "Admin/next-level approver completes final authorization based on hierarchy and policy.")
add_bullet(doc, "Store/issuance processing team issues approved items.")
add_bullet(doc, "System updates stock balances and preserves audit history for all actions.")

add_heading(doc, "3. Procurement to Stock Flow (As Implemented)", 12)
add_bullet(doc, "Create and manage tenders (Contract / Annual / Spot Purchase).")
add_bullet(doc, "Manage vendor participation, evaluation, and award decisions.")
add_bullet(doc, "Generate purchase orders from approved procurement outcomes.")
add_bullet(doc, "Receive delivery against PO with quantity and quality checks (including partial deliveries).")
add_bullet(doc, "Post received quantities into inventory and reflect updates in stock views.")

add_heading(doc, "4. Month-Wise Progress Snapshot", 12)
add_bullet(doc, "Month 1: Requirements finalization, architecture planning, database foundation.")
add_bullet(doc, "Month 2: Master data and core inventory structures implemented.")
add_bullet(doc, "Month 3: Tender, vendor, purchase-order lifecycle implemented.")
add_bullet(doc, "Month 4: Stock issuance and multi-level approval workflows implemented.")
add_bullet(doc, "Month 5: Store keeper verification workflow, audit trail, and control hardening completed.")
add_bullet(doc, "Month 6: End-to-end integration, stabilization, SQL enhancements, and release readiness completed.")

add_heading(doc, "5. Key Deliverables Completed", 12)
add_bullet(doc, "End-to-end stock request-to-issuance workflow.")
add_bullet(doc, "Tender-to-PO-to-delivery procurement chain.")
add_bullet(doc, "Store keeper physical verification integrated with approval flow.")
add_bullet(doc, "Role-based access and permission governance.")
add_bullet(doc, "Audit trail, transaction-safe processing, and soft-delete data safety.")
add_bullet(doc, "Financial-year aware inventory reporting support.")

add_heading(doc, "6. Current Status and Next Focus", 12)
doc.add_paragraph(
    "Core IMS modules are complete and operational for business use. The main next focus area is advanced reporting and "
    "analytics dashboards for management insights and KPI-driven monitoring."
)

output_file = "IMS-6-Month-Progress-One-Pager-Proper-Flow.docx"
doc.save(output_file)
print(f"Created: {output_file}")
