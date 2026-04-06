from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


doc = Document()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INVENTORY MANAGEMENT SYSTEM (IMS)\nTOTAL SYSTEM OVERVIEW (ONE-PAGER)")
run.bold = True
run.font.size = Pt(16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("System Summary for Project Management Review").font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "1. System Purpose")
doc.add_paragraph(
    "The Inventory Management System (IMS) is an enterprise web platform that manages the full lifecycle of inventory: "
    "planning and procurement, purchase ordering, delivery receiving, stock control, request approvals, issuance, and "
    "verification. It provides role-based access, auditability, and multi-level operational control for organization-wide usage."
)

add_heading(doc, "2. Core Modules")
add_bullet(doc, "Procurement Management: Contract, Annual, and Spot/Patty tender workflows.")
add_bullet(doc, "Vendor Management: Vendor registration, participation, assignment, and evaluation.")
add_bullet(doc, "Purchase Order Management: PO creation, line-item tracking, and PO lifecycle control.")
add_bullet(doc, "Delivery Receiving: PO-linked receiving, quality notes, partial delivery handling, and stock posting.")
add_bullet(doc, "Stock Issuance: Item request, approval routing, processing, and final issuance.")
add_bullet(doc, "Approval Workflow: Multi-level approval with item-wise decisions and forwarding.")
add_bullet(doc, "Store Keeper Verification: Physical stock confirmation with available/partial/unavailable outcome.")
add_bullet(doc, "Inventory Control: Current inventory visibility, stock movements, year-wise balance support.")
add_bullet(doc, "Administration & Security: User roles, permission governance, and operational audit trail.")

add_heading(doc, "3. End-to-End Business Flow")
add_bullet(doc, "Procurement Flow: Tender Creation -> Vendor Selection -> PO Generation -> Delivery Receipt -> Stock Update.")
add_bullet(doc, "Issuance Flow: User Request -> Supervisor Review -> Verification (if needed) -> Final Approval -> Issue -> Stock Deduction.")
add_bullet(doc, "Verification Flow: Approver Forwards Item -> Store Keeper Physical Check -> Status Feedback -> Approval Finalization.")

add_heading(doc, "4. Inventory Model")
add_bullet(doc, "Three-level inventory structure implemented: Admin Store, Wing Store, and Personal allocation level.")
add_bullet(doc, "Real-time status visibility and transaction-driven stock updates across levels.")
add_bullet(doc, "Financial-year aware inventory reporting support through year-wise inventory views.")

add_heading(doc, "5. Roles and Governance")
add_bullet(doc, "General Users: Raise item requests and track status.")
add_bullet(doc, "Wing Supervisors/Approvers: Review and route requests, trigger verification workflows.")
add_bullet(doc, "Store Keepers: Perform physical verification and confirm stock availability.")
add_bullet(doc, "Admins/Management: Final approval authority, configuration, monitoring, and oversight.")
add_bullet(doc, "Role-based permissions and workflow history provide accountability and compliance support.")

add_heading(doc, "6. Technical Overview")
add_bullet(doc, "Frontend: React + TypeScript + Vite with dashboard-based module navigation.")
add_bullet(doc, "Backend: Node.js/Express APIs for procurement, inventory, approvals, users, and workflows.")
add_bullet(doc, "Database: SQL Server with workflow procedures, views, and migration scripts.")
add_bullet(doc, "Data Integrity: Transaction-safe operations, status-driven workflows, and soft-delete strategy.")

add_heading(doc, "7. Current System Status")
doc.add_paragraph(
    "IMS core modules are integrated and operational for production use. The system currently supports end-to-end inventory "
    "operations with approval controls, verification traceability, and procurement-to-stock continuity."
)

output_file = "IMS-Total-System-Overview-One-Pager.docx"
doc.save(output_file)
print(f"Created: {output_file}")
